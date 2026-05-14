"""
SkillBaseHire — Resume Parser
Extracts structured data from PDF / DOCX / DOC files using section detection
and regex-based heuristics. All output is for candidate review before saving.
"""

import re

# ── Section aliases ───────────────────────────────────────────────────────────

_SECTION_MAP = {
    'summary': [
        'summary', 'objective', 'professional summary', 'career objective',
        'profile', 'about me', 'overview', 'professional profile', 'about',
        'personal statement', 'career summary',
    ],
    'experience': [
        'experience', 'work experience', 'employment', 'employment history',
        'professional experience', 'work history', 'career history',
        'professional background', 'internships', 'internship experience',
        'work & experience',
    ],
    'education': [
        'education', 'academic background', 'educational background',
        'academic qualifications', 'qualifications', 'educational qualifications',
        'academic history', 'academics',
    ],
    'skills': [
        'skills', 'technical skills', 'core skills', 'key skills', 'competencies',
        'expertise', 'technologies', 'technical expertise', 'core competencies',
        'programming languages', 'tools & technologies', 'tools and technologies',
        'tools', 'languages', 'frameworks', 'software skills', 'technical stack',
        'technology stack', 'skill set',
    ],
    'projects': [
        'projects', 'personal projects', 'academic projects', 'portfolio',
        'key projects', 'project work', 'notable projects', 'major projects',
        'project details',
    ],
    'certifications': [
        'certifications', 'certificates', 'professional certifications',
        'achievements', 'awards', 'awards & certifications',
        'awards and certifications', 'training', 'courses', 'online courses',
        'licenses & certifications', 'licenses and certifications',
        'professional development',
    ],
}

# Build reverse lookup once
_ALIAS_LOOKUP: dict[str, str] = {}
for _sec, _aliases in _SECTION_MAP.items():
    for _a in _aliases:
        _ALIAS_LOOKUP[_a.lower()] = _sec

# ── Date patterns ─────────────────────────────────────────────────────────────

_MONTH_PAT = (
    r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
)
_YEAR_PAT  = r'(?:19|20)\d{2}'

_DATE_TOKEN_RE = re.compile(
    rf'(?:{_MONTH_PAT}\.?\s*{_YEAR_PAT}|{_YEAR_PAT})',
    re.IGNORECASE,
)
_DATE_RANGE_RE = re.compile(
    rf'(?:{_MONTH_PAT}\.?\s*{_YEAR_PAT}|{_YEAR_PAT})'
    r'\s*[-–—to/]+\s*'
    rf'(?:{_MONTH_PAT}\.?\s*{_YEAR_PAT}|{_YEAR_PAT}'
    r'|[Pp]resent|[Cc]urrent|[Tt]ill\s+[Dd]ate|[Nn]ow|[Oo]ngoing)',
    re.IGNORECASE,
)
_YEAR_RE = re.compile(rf'\b({_YEAR_PAT})\b')
_EXP_YEARS_RE = re.compile(
    r'(\d+(?:\.\d+)?)\+?\s*(?:year|yr)s?\s*(?:of\s+)?(?:experience|exp(?:erience)?)',
    re.IGNORECASE,
)

# ── Degree keywords ───────────────────────────────────────────────────────────

_DEGREE_RE = re.compile(
    r'\b(?:B\.?\s*Tech|B\.?\s*E\.?\b|B\.?\s*Sc\.?\b|B\.?\s*C\.?\s*A\.?\b|'
    r'B\.?\s*Com\.?\b|B\.?\s*A\.?\b|M\.?\s*Tech|M\.?\s*E\.?\b|M\.?\s*Sc\.?\b|'
    r'M\.?\s*C\.?\s*A\.?\b|M\.?\s*Com\.?\b|M\.?\s*B\.?\s*A\.?\b|M\.?\s*A\.?\b|'
    r'Ph\.?\s*D\.?\b|Bachelor\'?s?|Master\'?s?|Diploma|Intermediate|'
    r'High\s+School|Secondary|10th|12th|SSC|HSC|Bachelors?|Masters?)\b',
    re.IGNORECASE,
)

# ── Known skills for fallback matching ───────────────────────────────────────

_KNOWN_SKILLS = {s.lower(): s for s in [
    'Python', 'JavaScript', 'TypeScript', 'Java', 'C++', 'C#', 'C', 'Go',
    'Rust', 'Ruby', 'PHP', 'Swift', 'Kotlin', 'Scala', 'R',
    'React', 'Angular', 'Vue', 'Vue.js', 'Next.js', 'Nuxt.js', 'Svelte',
    'Node.js', 'Express', 'Express.js', 'Django', 'Flask', 'FastAPI',
    'Spring', 'Spring Boot', 'Laravel', 'Rails', 'ASP.NET',
    'HTML', 'CSS', 'HTML/CSS', 'Tailwind', 'Bootstrap', 'Sass', 'SCSS',
    'SQL', 'PostgreSQL', 'MySQL', 'SQLite', 'MongoDB', 'Redis', 'DynamoDB',
    'Cassandra', 'Elasticsearch', 'Oracle',
    'AWS', 'Azure', 'GCP', 'Firebase', 'Supabase', 'Heroku',
    'Docker', 'Kubernetes', 'Terraform', 'Ansible', 'Jenkins',
    'Git', 'GitHub', 'GitLab', 'Linux', 'Unix', 'Bash',
    'CI/CD', 'DevOps', 'Agile', 'Scrum', 'JIRA',
    'Machine Learning', 'Deep Learning', 'NLP', 'TensorFlow', 'PyTorch',
    'Pandas', 'NumPy', 'Scikit-learn', 'OpenCV', 'Keras',
    'REST', 'REST APIs', 'GraphQL', 'gRPC', 'Microservices',
    'Selenium', 'Jest', 'PyTest', 'Cypress', 'Playwright', 'Mocha',
    'Figma', 'Webpack', 'Vite', 'Redux', 'Zustand',
]}

_DESIGNATION_RE = re.compile(
    r'\b(?:engineer|developer|analyst|manager|director|architect|designer|'
    r'consultant|lead|senior|junior|associate|intern|specialist|scientist|'
    r'programmer|administrator|coordinator|officer|head|vp|cto|ceo|coo|'
    r'executive|president|founder|co-founder|trainee|fresher)\b',
    re.IGNORECASE,
)

_CONTACT_RE = re.compile(
    r'@|www\.|linkedin|github|http|\+\d|\d{10}|email|phone|mobile',
    re.IGNORECASE,
)

_MONTH_MAP = {
    'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04',
    'may': '05', 'jun': '06', 'jul': '07', 'aug': '08',
    'sep': '09', 'oct': '10', 'nov': '11', 'dec': '12',
}


# ── Public API ────────────────────────────────────────────────────────────────

def extract_text(filepath: str, ext: str) -> str:
    """Extract plain text from a PDF, DOCX, or DOC file."""
    ext = ext.lower().lstrip('.')
    try:
        if ext == 'pdf':
            return _extract_pdf(filepath)
        if ext == 'docx':
            return _extract_docx(filepath)
        if ext == 'doc':
            return _extract_doc(filepath)
    except Exception:
        pass
    return ''


def parse_resume(text: str) -> dict:
    """
    Parse a plain-text resume and return a structured dict:
    {
        headline, job_title, experience, location,
        skills:          [str, ...],
        work_experience: [{company, designation, start_date, end_date, is_current, description}, ...],
        education:       [{degree, college, start_year, end_year}, ...],
        projects:        [{project_name, domain, description, project_url, year}, ...],
        certifications:  [{cert_name, issued_by, year, credential_url}, ...],
    }
    """
    lines    = [l.rstrip() for l in text.split('\n')]
    sections = _detect_sections(lines)
    return {
        'headline':        _extract_headline(lines, sections),
        'job_title':       _extract_job_title(lines, sections),
        'experience':      _extract_exp_years(text, sections),
        'location':        _extract_location(text),
        'skills':          _extract_skills(sections, text),
        'work_experience': _extract_work_experience(sections),
        'education':       _extract_education(sections),
        'projects':        _extract_projects(sections),
        'certifications':  _extract_certifications(sections),
    }


# ── Text extractors ───────────────────────────────────────────────────────────

def _extract_pdf(filepath: str) -> str:
    from pdfminer.high_level import extract_text as _pdfmine
    text = _pdfmine(filepath) or ''
    # Normalise ligatures and whitespace
    text = text.replace('\x0c', '\n')
    return text


def _extract_docx(filepath: str) -> str:
    from docx import Document
    doc   = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return '\n'.join(parts)


def _extract_doc(filepath: str) -> str:
    try:
        import docx2txt
        return docx2txt.process(filepath) or ''
    except ImportError:
        return ''


# ── Section detection ─────────────────────────────────────────────────────────

def _detect_sections(lines: list) -> dict:
    """
    Walk lines top-to-bottom.  When a short line matches a known section
    header alias, switch the accumulator to that section.
    Lines before the first recognised header go into 'header'.
    """
    sections: dict = {'header': []}
    current = 'header'

    for line in lines:
        norm = line.strip().lower().rstrip(':').rstrip('.').strip()
        # Header candidate: 1–6 words, no digits-heavy content
        if norm and 1 <= len(norm.split()) <= 6 and norm in _ALIAS_LOOKUP:
            current = _ALIAS_LOOKUP[norm]
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, [])
            sections[current].append(line)

    return sections


# ── Field extractors ──────────────────────────────────────────────────────────

def _extract_headline(lines: list, sections: dict) -> str:
    header_lines = [l.strip() for l in sections.get('header', []) if l.strip()]
    # Skip line 0 (usually the candidate's name) and contact lines
    for line in header_lines[1:5]:
        if not _CONTACT_RE.search(line) and len(line) > 4:
            return line[:120]
    # Fallback: first line of summary
    for line in sections.get('summary', []):
        l = line.strip()
        if l:
            return l[:120]
    return ''


def _extract_job_title(lines: list, sections: dict) -> str:
    # Same heuristic as headline — designation near the top of the resume
    header_lines = [l.strip() for l in sections.get('header', []) if l.strip()]
    for line in header_lines[1:5]:
        if not _CONTACT_RE.search(line) and _DESIGNATION_RE.search(line):
            return line[:120]
    return _extract_headline(lines, sections)


def _extract_exp_years(text: str, sections: dict) -> str:
    m = _EXP_YEARS_RE.search(text)
    if m:
        val = m.group(1)
        return str(int(float(val))) if '.' in val else val
    return ''


def _extract_location(text: str) -> str:
    header = text[:800]
    cities = re.findall(
        r'\b(?:Bangalore|Bengaluru|Mumbai|Delhi|New\s*Delhi|Hyderabad|Chennai|'
        r'Pune|Kolkata|Noida|Gurugram|Gurgaon|Ahmedabad|Jaipur|Surat|Kochi|'
        r'Chandigarh|Coimbatore|Indore|Bhopal|Nagpur|Vadodara|Lucknow|'
        r'New\s+York|San\s+Francisco|London|Sydney|Singapore|Remote)\b',
        header, re.IGNORECASE,
    )
    if cities:
        return cities[0].strip()
    m = re.search(
        r'\b([A-Z][a-z]{2,}(?:\s+[A-Z][a-z]{2,})?),\s*([A-Z][a-z]{2,}(?:\s+[A-Z][a-z]{2,})?)\b',
        header,
    )
    if m:
        return m.group(0)
    return ''


def _extract_skills(sections: dict, full_text: str) -> list:
    skills_lines = sections.get('skills', [])
    if not skills_lines:
        return _match_known_skills(full_text)

    raw_text = ' '.join(skills_lines)
    # Split on common delimiters
    raw_items = re.split(r'[,|•·\n\t/;]+', raw_text)
    cleaned = []
    seen = set()
    for item in raw_items:
        item = item.strip().strip('•·-–—○●▪▫▸►⮞*').strip()
        if item and 2 <= len(item) <= 50 and item.lower() not in seen:
            cleaned.append(item)
            seen.add(item.lower())

    return cleaned[:60] if cleaned else _match_known_skills(raw_text)


def _match_known_skills(text: str) -> list:
    found = []
    seen  = set()
    for token_lower, canonical in _KNOWN_SKILLS.items():
        if re.search(r'\b' + re.escape(token_lower) + r'\b', text, re.IGNORECASE):
            if token_lower not in seen:
                found.append(canonical)
                seen.add(token_lower)
    return found


def _extract_work_experience(sections: dict) -> list:
    exp_lines = sections.get('experience', [])
    if not exp_lines:
        return []

    entries  = []
    current  = None
    desc_buf: list = []

    def _flush():
        if current is not None:
            current['description'] = ' '.join(desc_buf).strip()[:800]
            entries.append(current)

    for line in exp_lines:
        stripped = line.strip()
        if not stripped:
            continue

        date_m = _DATE_RANGE_RE.search(stripped)
        if date_m:
            _flush()
            current  = None
            desc_buf = []

            date_str   = date_m.group(0)
            is_current = bool(re.search(
                r'[Pp]resent|[Cc]urrent|[Tt]ill\s+[Dd]ate|[Nn]ow|[Oo]ngoing',
                date_str,
            ))
            tokens     = _DATE_TOKEN_RE.findall(date_str)
            start_date = _normalise_date(tokens[0]) if tokens else ''
            end_date   = '' if is_current else (_normalise_date(tokens[1]) if len(tokens) > 1 else '')

            remainder = (stripped[:date_m.start()] + stripped[date_m.end():]).strip()
            company, designation = _split_company_designation(remainder)

            current = dict(
                company     = company[:200],
                designation = designation[:200],
                start_date  = start_date,
                end_date    = end_date,
                is_current  = 1 if is_current else 0,
            )
        elif current is not None:
            clean = stripped.lstrip('•·-–—○●▪▫▸►*').strip()
            if clean:
                desc_buf.append(clean)
        elif _looks_like_entry_header(stripped):
            # Company or role line without an inline date range
            _flush()
            current  = dict(company=stripped[:200], designation='',
                            start_date='', end_date='', is_current=0)
            desc_buf = []

    _flush()
    return [e for e in entries if e.get('company') or e.get('designation')][:10]


def _extract_education(sections: dict) -> list:
    edu_lines = sections.get('education', [])
    if not edu_lines:
        return []

    entries = []
    current = None

    for line in edu_lines:
        stripped = line.strip()
        if not stripped:
            continue

        if _DEGREE_RE.search(stripped):
            if current:
                entries.append(current)
            years      = _YEAR_RE.findall(stripped)
            start_year = years[0]  if years else ''
            end_year   = years[-1] if len(years) > 1 else ''
            remainder  = _YEAR_RE.sub('', stripped).strip().strip('-–—to').strip()
            degree, college = _split_degree_college(remainder)
            current = dict(degree=degree[:200], college=college[:200],
                           start_year=start_year, end_year=end_year)
        elif current and not current.get('college'):
            # College name may be on the next line after the degree line
            if not _YEAR_RE.search(stripped) and not _DEGREE_RE.search(stripped):
                current['college'] = stripped[:200]

    if current:
        entries.append(current)

    return [e for e in entries if e.get('degree')][:6]


def _extract_projects(sections: dict) -> list:
    proj_lines = sections.get('projects', [])
    if not proj_lines:
        return []

    entries  = []
    current  = None
    desc_buf: list = []

    def _flush():
        if current is not None:
            current['description'] = ' '.join(desc_buf).strip()[:600]
            entries.append(current)

    for line in proj_lines:
        stripped = line.strip()
        if not stripped:
            continue

        is_header = (
            not stripped[0] in '•·-–—○●▪▫▸►*'
            and len(stripped.split()) <= 10
            and stripped[0].isupper()
            and not _DATE_RANGE_RE.search(stripped)
            and not _CONTACT_RE.search(stripped)
        )
        if is_header:
            _flush()
            years = _YEAR_RE.findall(stripped)
            year  = years[-1] if years else ''
            name  = _YEAR_RE.sub('', stripped).strip().strip('|-–—:').strip()
            current  = dict(project_name=(name or stripped)[:200],
                            domain='', description='', project_url='', year=year)
            desc_buf = []
        elif current is not None:
            clean = stripped.lstrip('•·-–—○●▪▫▸►*').strip()
            url_m = re.search(r'https?://\S+', clean)
            if url_m and not current['project_url']:
                current['project_url'] = url_m.group(0)[:300]
            if clean:
                desc_buf.append(clean)

    _flush()
    return [e for e in entries if e.get('project_name')][:8]


def _extract_certifications(sections: dict) -> list:
    cert_lines = sections.get('certifications', [])
    if not cert_lines:
        return []

    entries = []
    for line in cert_lines:
        stripped = line.strip().lstrip('•·-–—○●▪▫▸►*').strip()
        if not stripped or len(stripped) < 4:
            continue

        years = _YEAR_RE.findall(stripped)
        year  = years[-1] if years else ''
        name  = _YEAR_RE.sub('', stripped).strip().strip('|-–—():').strip()
        if not name:
            continue

        issued_by = ''
        for sep in (' by ', ' from ', ' – ', ' - ', ' | ', ' — '):
            if sep.lower() in name.lower():
                idx   = name.lower().index(sep.lower())
                issued_by = name[idx + len(sep):].strip()
                name      = name[:idx].strip()
                break
        paren_m = re.match(r'^(.+?)\s*\(([^)]+)\)\s*$', name)
        if paren_m:
            name, issued_by = paren_m.group(1).strip(), paren_m.group(2).strip()

        entries.append(dict(cert_name=name[:200], issued_by=issued_by[:200],
                            year=year, credential_url=''))

    return entries[:10]


# ── Helpers ───────────────────────────────────────────────────────────────────

def _normalise_date(s: str) -> str:
    s = s.strip()
    m = re.match(
        r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s*(\d{4})',
        s, re.IGNORECASE,
    )
    if m:
        mon = _MONTH_MAP.get(m.group(1).lower()[:3], '01')
        return f"{m.group(2)}-{mon}-01"
    m = re.match(r'(\d{4})', s)
    if m:
        return f"{m.group(1)}-01-01"
    return s


def _split_company_designation(text: str) -> tuple:
    for sep in (' at ', ' @ ', ' – ', ' — ', ' - ', ' | ', ', '):
        if sep in text:
            a, b = text.split(sep, 1)
            a, b = a.strip(), b.strip()
            if _DESIGNATION_RE.search(b):
                return a, b
            if _DESIGNATION_RE.search(a):
                return b, a
            return a, b
    return text.strip(), ''


def _split_degree_college(text: str) -> tuple:
    for sep in (' from ', ' at ', ' – ', ' - ', ' | ', ', '):
        if sep in text:
            parts = text.split(sep, 1)
            return parts[0].strip(), parts[1].strip()
    return text.strip(), ''


def _looks_like_entry_header(text: str) -> bool:
    if len(text) > 80 or text[0] in '•·-–—○●▪▫▸►*':
        return False
    return bool(re.match(r'^[A-Za-z0-9\s&.,\-()/]+$', text))
